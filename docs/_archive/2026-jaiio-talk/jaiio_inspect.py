"""Inspect the JAIIO .docx structure for editing."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

PATH = r'C:\Users\lsotomayor\Downloads\Graph_Hunter_JAIIO_ES.docx'
d = Document(PATH)

print(f'paragraphs: {len(d.paragraphs)}')
print(f'tables: {len(d.tables)}')
print()
print('=== paragraphs (idx, first 100 chars) ===')
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    style = p.style.name if p.style else '?'
    print(f'{i:3d} [{style[:14]:14s}] {t[:100]}')

print()
print('=== tables ===')
for ti, tbl in enumerate(d.tables):
    print(f'Table {ti}: {len(tbl.rows)} rows x {len(tbl.columns)} cols')
    for ri, row in enumerate(tbl.rows[:3]):
        cells = [c.text.strip()[:30] for c in row.cells]
        print(f'  row {ri}: {cells}')
