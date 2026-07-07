"""Restore bold prefixes lost during the main edit pass.

The first pass used a simple run-replacement strategy that dropped run-level
formatting. This script re-renders specific paragraphs with a leading bold
fragment followed by plain text, matching the paper's house style.
"""
from __future__ import annotations

import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

SRC = Path(r'C:\Users\lsotomayor\Downloads\Graph_Hunter_JAIIO_ES.docx')


def set_paragraph_runs(p, bold_prefix: str, plain_rest: str) -> None:
    """Replace a paragraph's runs with [bold_prefix][plain_rest]."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.bold = True
    if plain_rest:
        p.add_run(plain_rest)


def find_para_by_text(doc, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise LookupError(f'paragraph not found: {needle!r}')


def main() -> int:
    doc = Document(str(SRC))

    # §4.7 heading — match style of other 4.x headings (bold).
    p = find_para_by_text(
        doc,
        '4.7 Compilación asistida por LLM con muestreo restringido',
    )
    set_paragraph_runs(p, p.text, '')

    # §6.1 — restore "Fragmentación del contexto." bold prefix.
    p = find_para_by_text(doc, 'Fragmentación del contexto. Tres mecanismos')
    full = p.text
    prefix = 'Fragmentación del contexto.'
    rest = full[len(prefix):]
    set_paragraph_runs(p, prefix, rest)

    # §6.2 — restore "Escalabilidad en memoria." bold prefix.
    p = find_para_by_text(doc, 'Escalabilidad en memoria. El enfoque en memoria')
    full = p.text
    prefix = 'Escalabilidad en memoria.'
    rest = full[len(prefix):]
    set_paragraph_runs(p, prefix, rest)

    # §6.2 — bold prefix on the new "Dependencia del modelo local..." bullet.
    p = find_para_by_text(doc, 'Dependencia del modelo local en la vía LLM.')
    full = p.text
    prefix = 'Dependencia del modelo local en la vía LLM.'
    rest = full[len(prefix):]
    set_paragraph_runs(p, prefix, rest)

    # §7 — restore "Generación asistida de hipótesis." bold prefix with colon
    # to match the style of sibling bullets ("Búsqueda en streaming:", etc.).
    p = find_para_by_text(doc, 'Generación asistida de hipótesis.')
    full = p.text
    # Convert "Generación asistida de hipótesis." → "Generación asistida de hipótesis:"
    if 'Generación asistida de hipótesis.' in full:
        full = full.replace(
            'Generación asistida de hipótesis.',
            'Generación asistida de hipótesis:',
            1,
        )
    prefix = 'Generación asistida de hipótesis:'
    rest = full[len(prefix):]
    set_paragraph_runs(p, prefix, rest)

    doc.save(str(SRC))
    print(f'format fixes applied: {SRC}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
