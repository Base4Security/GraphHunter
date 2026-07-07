"""Rebuild the §7 'Generación asistida de hipótesis' bullet from scratch.

The previous format-fix step miscalculated the slice index for the new
bullet (the leading '• ' was not accounted for), producing the corrupted
text 'Generación asistida de hipótesis:s: La infraestructura...'.
This script overwrites that paragraph with the correct bullet structure,
matching the sibling bullets in §7.
"""
from __future__ import annotations

import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

SRC = Path(r'C:\Users\lsotomayor\Downloads\Graph_Hunter_JAIIO_ES.docx')


def main() -> int:
    doc = Document(str(SRC))
    target = None
    for p in doc.paragraphs:
        if 'Generación asistida de hipótesis' in p.text:
            target = p
            break
    if target is None:
        print('ERROR: target paragraph not found', file=sys.stderr)
        return 1

    for r in list(target.runs):
        r._element.getparent().remove(r._element)

    target.add_run('• ')
    bold = target.add_run('Generación asistida de hipótesis:')
    bold.bold = True
    target.add_run(
        ' La infraestructura de muestreo restringido por gramática '
        'introducida para la compilación de parsers se generaliza '
        'directamente a la producción de hipótesis: una gramática que '
        'describe la sintaxis de cadenas tipadas consistentes con la '
        'ontología (Σ_V, Σ_E) restringe al LLM a sugerir únicamente '
        'hipótesis bien-formadas, alimentadas por embeddings del grafo y '
        'scores GNN, complementando el catálogo ATT&CK existente.'
    )

    doc.save(str(SRC))
    print(f'wrote: {SRC}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
