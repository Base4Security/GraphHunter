"""Apply targeted insertions/edits to the JAIIO Graph Hunter paper (.docx).

Anchors paragraphs by substring match and either replaces text content of
existing paragraphs or inserts new paragraphs before a known anchor. Tables
are amended via add_row().

Run: python jaiio_update.py
"""
from __future__ import annotations

import shutil
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

SRC = Path(r'C:\Users\lsotomayor\Downloads\Graph_Hunter_JAIIO_ES.docx')
BAK = SRC.with_suffix('.docx.bak')


def find_para(doc, needle: str, *, start: int = 0) -> int:
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if needle in p.text:
            return i
    raise LookupError(f'anchor not found: {needle!r}')


def replace_para_text(p, new_text: str) -> None:
    """Replace a paragraph's textual content, preserving the paragraph style.

    Strategy: drop all existing runs, then add a fresh run with the new text.
    Run-level character formatting (bold, italic spans) is lost on edited
    paragraphs — acceptable here since we only edit body paragraphs that are
    single-formatted.
    """
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(new_text)


def insert_para_before(p_anchor, text: str, style=None):
    """Insert a new paragraph immediately before p_anchor, copying its style."""
    new_p = p_anchor.insert_paragraph_before(text)
    if style is None:
        new_p.style = p_anchor.style
    else:
        new_p.style = style
    return new_p


def insert_para_after(doc, idx: int, text: str):
    """Insert a paragraph after doc.paragraphs[idx], matching its style."""
    # Use the anchor at idx+1 (or fall back to end-of-doc) and insert before it.
    if idx + 1 < len(doc.paragraphs):
        return insert_para_before(doc.paragraphs[idx + 1], text)
    # No paragraph after — append at end via document body.
    p = doc.add_paragraph(text)
    p.style = doc.paragraphs[idx].style
    return p


def main() -> int:
    if not SRC.exists():
        print(f'ERROR: source missing: {SRC}', file=sys.stderr)
        return 1

    # Backup once (keep first backup; do not overwrite).
    if not BAK.exists():
        shutil.copy2(SRC, BAK)
        print(f'backup: {BAK}')
    else:
        print(f'backup already exists: {BAK} (kept)')

    doc = Document(str(SRC))

    # ---------- §1.3 Propuesta y objetivos ----------
    i = find_para(doc, '6. Integración de LLM mediante MCP')
    replace_para_text(
        doc.paragraphs[i],
        '6. Integrar LLM mediante un servidor Model Context Protocol (MCP) que '
        'exponga las operaciones del motor — organizadas por área funcional '
        '(catálogo de datasets, travesía del grafo, formulación y ejecución de '
        'hipótesis, ingesta, detalle de nodos, anotaciones, scoring) — como '
        '"herramientas" que un asistente de IA puede invocar, manteniendo al '
        'analista humano en el control de las decisiones.'
    )
    insert_para_after(
        doc, i,
        '7. Incorporar una vía de compilación asistida por modelo de lenguaje '
        '(LLM) local con muestreo restringido por gramática, capaz de '
        'recuperar mapeos de campos para fuentes de telemetría no '
        'estandarizadas donde el detector heurístico produce cero o pocos '
        'triples, sin requerir servicios externos ni intervención manual de '
        'configuración.'
    )

    # ---------- §3.2 Estructuras y mapeo de logs a grafos ----------
    i = find_para(doc, '• Sysmon: Procesa eventos de Windows Sysmon')
    replace_para_text(
        doc.paragraphs[i],
        '• Sysmon: Procesa eventos de Windows Sysmon (creación de procesos, '
        'conexiones de red, accesos a archivos, consultas DNS, modificaciones '
        'de registro, eventos WMI). Los EventIDs 1, 3, 5, 7, 8, 11, 12/13, 15 '
        'y 23 se decodifican por una vía tipada nativa basada en simd-json que '
        'evita los costos de un parser JSON dinámico genérico; el resto se '
        'enruta al parser JSON con mapeo por campo. Por ejemplo, un EventID 1 '
        '(Process Create) genera una tripleta (ParentProcess, Spawn, '
        'ChildProcess).'
    )

    i = find_para(doc, '• Azure Sentinel: Procesa logs del SIEM cloud')
    replace_para_text(
        doc.paragraphs[i],
        '• Azure Sentinel: Procesa logs del SIEM cloud de Microsoft con una '
        'vía nativa por tabla — SecurityEvent (eventos Windows reenviados), '
        'SigninLogs (autenticaciones Azure AD), AzureActivity (operaciones '
        'sobre recursos cloud), CommonSecurityLog (datos CEF) y las tablas de '
        'Microsoft Defender for Endpoint DeviceProcessEvents, '
        'DeviceNetworkEvents y DeviceFileEvents. Los EventIDs 4624/4625 '
        '(logon exitoso/fallido) generan tripletas (User, Auth, Host).'
    )

    # New bullet: Windows Security + PowerShell EVTX fast path.
    i = find_para(doc, '• CSV: Procesa datos tabulares')
    insert_para_after(
        doc, i,
        '• Windows Security y PowerShell: Para colectores que entregan logs '
        'EVTX nativos, el sistema decodifica tipadamente los EventIDs más '
        'frecuentes en el ciclo de Threat Hunting — 4624/4625 (logon '
        'exitoso/fallido), 4688 (creación de proceso), 4689 (terminación), '
        '4663 (acceso a objeto), 5145 (acceso a recurso compartido) y 5156 '
        '(Windows Filtering Platform) del canal Security, más el EventID '
        '4104 (Script Block Logging) del canal PowerShell. Cada EventID '
        'mapea a tripletas específicas sin pasar por la materialización '
        'intermedia en JSON.'
    )

    # New paragraph: RawIngestEvent as intermediate representation.
    i = find_para(doc, 'La política de ingesta emplea upsert')
    insert_para_after(
        doc, i,
        'Para evitar materializaciones intermedias en cadenas de texto, los '
        'parsers tipados producen registros RawIngestEvent — un struct con '
        'identificador y tipo de cada extremo, tipo de relación, marca '
        'temporal y diccionarios de metadatos — que se escriben en el grafo '
        'mediante insert_raw_events / insert_raw_events_chunk. Esta '
        'representación intermedia preserva el tipado de origen y permite '
        'lotear inserciones manteniendo la deduplicación de vértices y la '
        'estampa de dataset_id en cada arista para trazabilidad de '
        'procedencia.'
    )

    # ---------- §4.1 Visión general ----------
    # Insert adaptive-ingest paragraph after the in-memory rationale.
    i = find_para(doc, 'Además, el control explícito sobre el algoritmo')
    insert_para_after(
        doc, i,
        'La capa de ingestión es adaptativa: la previa del archivo enriquece '
        'cada columna detectada con muestras de valores, una sugerencia de '
        'tipo canónico (canonical_target), conteo de ocurrencias no vacías '
        'y, para campos temporales, hints de formato y locale que se aplican '
        'antes de la normalización a ISO-8601. El schema efectivamente '
        'aplicado, junto con las estadísticas reales de parse (filas leídas, '
        'filas que produjeron triples, ocurrencias por campo) se persisten '
        'en la metadata del dataset y se exponen en la tarjeta del dataset y '
        'en una alerta dedicada cuando una ingesta no produce ningún triple '
        '(zero_triples), permitiendo al analista corregir el mapeo sin '
        'reabrir la sesión.'
    )

    # Insert §4.7 preview paragraph immediately after the adaptive-ingest one.
    # (The just-inserted para is now at i+1; place the next one at i+2.)
    new_idx = i + 1
    insert_para_after(
        doc, new_idx,
        'Para fuentes cuyo mapeo heurístico no produce triples, el sistema '
        'incluye una vía de compilación asistida por un LLM local con '
        'muestreo restringido por gramática (descrita en §4.7), invocada '
        'únicamente como fallback y activable sin configuración manual '
        'mediante un dispositor plug-and-play.'
    )

    # ---------- §4.2 Estructuras de datos ----------
    # Change "cinco" → "seis" in the header sentence.
    i = find_para(doc, 'El motor mantiene cinco estructuras de datos en memoria')
    replace_para_text(
        doc.paragraphs[i],
        'El motor mantiene seis estructuras de datos en memoria:'
    )

    # Append a 6th bullet about RawIngestEvent path and a closing sentence
    # about HuntResult shape, inserted after item 5.
    i = find_para(doc, '5. Índice secundario de relaciones')
    insert_para_after(
        doc, i,
        '6. Pool de eventos crudos (Vec<RawIngestEvent>): buffer transitorio '
        'que recibe los registros tipados producidos por los parsers nativos '
        '(Sysmon, Windows Security, PowerShell, tablas Sentinel cloud) y los '
        'vuelca en lotes al grafo mediante insert_raw_events_chunk, '
        'preservando tipos y estampando dataset_id en cada arista sin pasar '
        'por la representación intermedia tripleta-string.'
    )
    # And right after that bullet, a closing remark on HuntResult shape.
    insert_para_after(
        doc, i + 1,
        'Como decisión deliberada de hot-loop, el tipo HuntResult se define '
        'como Vec<Arc<str>> sobre el interner de strings del grafo: cada '
        'paso emitido por la búsqueda incurre en un incremento de refcount en '
        'lugar de una asignación de String, eliminando la presión sobre el '
        'allocator cuando la consulta retorna miles de caminos.'
    )

    # ---------- §4.3 DFS con poda por dominio ----------
    # Update the preamble line about the algorithm's runtime/implementation.
    i = find_para(doc, 'El algoritmo central, Temporal-Hypothesis-Search')
    replace_para_text(
        doc.paragraphs[i],
        'El algoritmo central, Temporal-Hypothesis-Search, está implementado '
        'íntegramente en Rust nativo (sin frontera FFI hacia bibliotecas '
        'externas) y opera en dos fases:'
    )

    # After "Tratabilidad práctica" paragraph, add Yannakakis closure paragraph.
    i = find_para(doc, 'Tratabilidad práctica:')
    insert_para_after(
        doc, i,
        'Sobre esta base, el sistema integra un planificador de consulta '
        'tipado inspirado en el algoritmo de Yannakakis: para cada hipótesis '
        'se precomputa un árbol de bags que materializa los cierres por tipo '
        'de relación y entidad, eliminando candidatos antes de la fase DFS. '
        'Tanto el bag tree como los snapshots de prefijos parciales se '
        'mantienen en una caché por sesión, invalidada únicamente cuando la '
        'versión de mutación del grafo cambia. Esto convierte hunts '
        'sucesivos sobre la misma sesión — el patrón habitual de '
        'investigación iterativa, donde el analista refina la hipótesis tras '
        'observar resultados parciales — en consultas dominadas por la fase '
        'de semillas, con un costo prácticamente constante frente al tamaño '
        'del grafo subyacente.'
    )

    # ---------- §4.4 Scoring y priorización ----------
    i = find_para(
        doc,
        'El score compuesto S = w₁·deg + w₂·PR + w₃·BC'
    )
    insert_para_after(
        doc, i,
        'El modo de scoring es seleccionable por consulta. Cada llamada a la '
        'API de ejecución de hipótesis acepta un override entre tres modos: '
        'Structural (DFS plano, sin penalización por anomalía), Anomaly '
        '(DFS guiado por los scores endógenos de §4.5, que también permite '
        'descartar tempranamente vecindarios con anomaly score bajo) y Gnn '
        '(idéntico a Anomaly pero exponiendo además la dimensión GNN en el '
        'desglose del resultado). El analista puede contrastar los tres '
        'modos sobre el mismo grafo sin recomputarlo, lo que facilita '
        'separar señales estructurales de señales aprendidas durante la '
        'investigación.'
    )

    # ---------- §4.6 MCP ----------
    i = find_para(doc, 'El sistema expone una API HTTP REST en localhost')
    replace_para_text(
        doc.paragraphs[i],
        'El sistema expone una API HTTP REST en localhost con autenticación '
        'Bearer token, permitiendo que herramientas externas consulten el '
        'grafo sin usar la interfaz gráfica. Sobre esta API, un servidor MCP '
        '(Model Context Protocol) traduce las operaciones del motor en '
        'herramientas invocables por asistentes de IA, organizadas en '
        'espacios por área funcional: catalog (descubrimiento de datasets), '
        'graph (lookup y travesía), hunt (formulación y ejecución de '
        'hipótesis con ranking por anomalía), ingest (detección de formato '
        'y carga), node (detalle de entidades, scores GNN, vecindario), '
        'notes (anotaciones de análisis), review (workflow de revisión), '
        'scores (operaciones de scoring por nodo y por camino), session '
        '(gestión del estado de la sesión activa) y tags (etiquetado de '
        'entidades y relaciones). Esta separación permite que un asistente '
        'de IA delegue tareas exploratorias acotadas sin tener acceso al '
        'conjunto completo de primitivas, manteniendo al humano en el loop '
        'para las decisiones de alto nivel.'
    )

    # ---------- §4.7 NEW: Compilación asistida por LLM ----------
    # Insert before "5. Evaluación".
    i_eval = find_para(doc, '5. Evaluación')
    p_eval = doc.paragraphs[i_eval]
    # Heading
    insert_para_before(
        p_eval,
        '4.7 Compilación asistida por LLM con muestreo restringido'
    )
    # Three body paragraphs.
    insert_para_before(
        p_eval,
        'Para fuentes no estandarizadas — exports CSV de plataformas '
        'internas, syslogs propietarios, registros tabulares de '
        'herramientas SaaS — el detector heurístico de campos puede '
        'producir un mapeo débil: ningún campo identificado como entidad y, '
        'en consecuencia, cero triples ingestados. En lugar de devolver el '
        'problema al analista, el sistema invoca como fallback un modelo de '
        'lenguaje local — Phi-3-mini cuantizado a 4 bits en formato GGUF, '
        'ejecutado por un backend de inferencia basado en Candle — para '
        'proponer un mapeo refinado. El modelo se ejecuta enteramente en '
        'CPU del cliente, sin red ni APIs externas, y se carga de forma '
        'perezosa la primera vez que la vía de fallback se activa.'
    )
    insert_para_before(
        p_eval,
        'La generación se realiza con muestreo restringido por gramática. '
        'En cada paso de decodificación, una máscara producida por un '
        'sampler que compila una gramática Lark (vrl_subset.lark) se '
        'aplica sobre los logits antes del muestreo, dejando vivos sólo los '
        'tokens admitidos por la gramática en el estado actual. La '
        'gramática define exactamente el subconjunto de VRL/JSON que el '
        'validador y el compilador de mapeos de campos aceptan, de modo '
        'que cualquier secuencia generada por el modelo es válida por '
        'construcción: el contrato del sistema es "gramática = validador = '
        'salida del compilador". Tres pruebas — compiler_roundtrip, '
        'fuzz_validator y grammar_roundtrip — pinean este invariante: si '
        'un cambio rompe la equivalencia entre las tres caras, falla la '
        'suite.'
    )
    insert_para_before(
        p_eval,
        'La vía LLM no reemplaza la heurística: la complementa. Un '
        'dispatcher ejecuta primero el detector heurístico y, sólo si éste '
        'no produce ningún mapeo Node, consulta el modelo, parsea su '
        'propuesta como un FieldConfig y la confronta con la heurística '
        'campo a campo. La diferencia se emite como evento al frontend, que '
        'muestra un banner side-by-side con ambas configuraciones, su '
        'confidence y el backend que la produjo; el analista decide cuál '
        'aplicar antes de re-ingestar. La activación es plug-and-play: el '
        'modelo se empaqueta como recurso del instalador y se resuelve por '
        'una cadena variable de entorno → bundle → directorio de trabajo, '
        'de modo que el badge de la aplicación refleja en tres estados — '
        'Ingest: v2, Ingest: v2 + LLM (idle) y Ingest: v2 + IA Compiler — '
        'si la vía está disponible, disponible-pero-inactiva o activa, sin '
        'requerir configuración manual.'
    )

    # ---------- §5.1 Entorno experimental ----------
    i = find_para(doc, 'Todas las mediciones utilizan compilaciones en modo release')
    replace_para_text(
        doc.paragraphs[i],
        'Todas las mediciones utilizan compilaciones en modo release de '
        'Rust sobre un solo core (AMD Ryzen 5 PRO 5650U, 16 GB RAM). El '
        'motor de matching corre en Rust nativo, sin frontera FFI, y no se '
        'emplea paralelismo en la fase de búsqueda para obtener mediciones '
        'reproducibles. Para el escenario de un millón de vértices se '
        'utilizan grafos sintéticos Barabási-Albert (V = 10⁶, m = 3, '
        '≈6 × 10⁶ aristas) que reproducen la distribución power-law '
        'observada en el corpus Sysmon real.'
    )

    # ---------- §5.5 Tabla 1 — append BA-Sysmon V=1M row ----------
    t = doc.tables[0]
    # Replicate the style of the last data row when adding a new one.
    template_row = t.rows[-1]
    new_row = t.add_row()
    values = [
        'BA-Sysmon',
        '1,000,000',
        '~6,000,000',
        '6',
        '478 s',
        '—',
        '1,000,000 †',
    ]
    for cell, value in zip(new_row.cells, values):
        cell.text = value
    # Copy paragraph alignment / run formatting from a template cell when possible.
    for new_cell, tmpl_cell in zip(new_row.cells, template_row.cells):
        if not tmpl_cell.paragraphs or not tmpl_cell.paragraphs[0].runs:
            continue
        tmpl_run = tmpl_cell.paragraphs[0].runs[0]
        tmpl_align = tmpl_cell.paragraphs[0].alignment
        # Mirror font name/size if set.
        for p in new_cell.paragraphs:
            if tmpl_align is not None:
                p.alignment = tmpl_align
            for r in p.runs:
                if tmpl_run.font.name:
                    r.font.name = tmpl_run.font.name
                if tmpl_run.font.size:
                    r.font.size = tmpl_run.font.size

    # Update post-table commentary: append a sentence about V=1M + ingest throughput.
    i = find_para(doc, 'El dataset Mordor, con 963K eventos originales')
    replace_para_text(
        doc.paragraphs[i],
        'El dataset Mordor, con 963K eventos originales comprimidos a 772 '
        'nodos y 21,963 aristas, presenta el caso real más exigente por su '
        'distribución power-law (d_max = 9,411). Aun así, la búsqueda con '
        'L=2 completa en 4.3 ms visitando sólo 13 nodos frente a 22,735 del '
        'baseline sin poda. La última fila de la tabla extiende la '
        'evaluación a escala sintética: sobre un grafo Barabási-Albert de '
        '10⁶ vértices y ≈6 × 10⁶ aristas, la enumeración exhaustiva de '
        '6-ciclos (L=6, cap = 10⁶ resultados) completa en una mediana de '
        '478 s sobre el mismo core — un orden de magnitud por debajo de '
        'los tiempos esperables para una consulta equivalente en un SIEM '
        'sobre dataset de telemetría comparable. El símbolo † indica que '
        'el cap configurable de resultados se saturó, lo que acota '
        'simultáneamente memoria y tiempo de respuesta.'
    )
    # Insert a follow-up paragraph about ingest throughput right after.
    insert_para_after(
        doc, i,
        'En el lado de la ingestión, la vía tipada con RawIngestEvent '
        'alcanza el techo de parse-only de ≈10⁶ eventos por segundo sobre '
        'un fixture mixto de Windows Security y PowerShell, y supera por '
        'un factor de ≈4,75× a la ruta JSON genérica sobre JSONL '
        'equivalente. Este caudal mantiene la ingestión por debajo del '
        'tiempo de la búsqueda en hunts interactivos, evitando que el '
        'parser se convierta en el cuello de botella aun en sesiones que '
        'cargan decenas de millones de eventos.'
    )

    # ---------- §6.1 Fragmentación del contexto ----------
    i = find_para(doc, 'Fragmentación del contexto. Tres mecanismos')
    p = doc.paragraphs[i]
    old = p.text
    # Append clause about adaptive ingest + LLM banner.
    addition = (
        ' La previa enriquecida y el banner de propuesta LLM extienden '
        'estos mecanismos a fuentes no estandarizadas: el analista resuelve '
        'mapeos débiles desde la misma sesión, sin abandonar la unificación '
        'bajo la ontología de nueve tipos de entidades y nueve tipos de '
        'relaciones.'
    )
    replace_para_text(p, old + addition)

    # ---------- §6.2 Limitaciones ----------
    i = find_para(doc, 'Escalabilidad en memoria. El enfoque en memoria')
    p = doc.paragraphs[i]
    old = p.text
    addition = (
        ' En el extremo opuesto, el benchmark sintético sobre 10⁶ vértices '
        'demuestra que la enumeración exhaustiva con cap de resultados '
        'configurable mantiene el consumo de memoria acotado incluso a '
        'esa escala, aunque al precio de minutos por consulta — un régimen '
        'aceptable para análisis batch pero no para hunt interactivo.'
    )
    replace_para_text(p, old + addition)

    # Add new limitation bullet after "Dependencia de modelos GNN pre-entrenados".
    i = find_para(doc, 'Dependencia de modelos GNN pre-entrenados')
    insert_para_after(
        doc, i,
        'Dependencia del modelo local en la vía LLM. La compilación '
        'asistida por LLM requiere un GGUF Phi-3-mini-Q4 (≈2.4 GB) '
        'empaquetado en el instalador o resuelto vía '
        'GRAPHHUNTER_LOCAL_MODEL_PATH; entornos sin el modelo en disco '
        'caen automáticamente al detector heurístico, y la vía completa '
        'puede desactivarse exponiendo GRAPHHUNTER_LLM_INGEST=0. La '
        'latencia de inferencia en CPU acota esta vía a uso ocasional como '
        'fallback, no al hot-path de ingesta masiva.'
    )

    # ---------- §7 Conclusiones — append sixth result + tweak future line ----------
    i = find_para(doc, 'Los resultados principales son: (1) aceleraciones')
    replace_para_text(
        doc.paragraphs[i],
        'Los resultados principales son: (1) aceleraciones de 1,747× a '
        '11,600× sobre baseline sin poda en datasets reales; (2) tiempos '
        'de búsqueda sub-milisegundo en grafos con más de 20,000 aristas; '
        '(3) reducción de un orden de magnitud en el MTTI (de 8–20 minutos '
        'a 1–2 minutos para investigaciones de 4 pasos); (4) descubrimiento '
        'exhaustivo de caminos de ataque versus la exploración parcial del '
        'pivoteo manual; (5) integración de clasificación GNN sobre '
        'subgrafos k-hop como quinto componente del scoring de anomalía, '
        'habilitando un ciclo cerrado de detección-investigación con '
        'retroalimentación bidireccional entre scoring endógeno y '
        'clasificación aprendida; y (6) extensión del régimen tratable a '
        'grafos de un millón de vértices — con enumeración exhaustiva de '
        '6-ciclos en escala de minutos — apoyada en una vía de ingestión '
        'tipada que sostiene ≈10⁶ eventos por segundo.'
    )

    i = find_para(doc, '• Generación automática de hipótesis')
    replace_para_text(
        doc.paragraphs[i],
        '• Generación asistida de hipótesis. La infraestructura de muestreo '
        'restringido por gramática introducida para la compilación de '
        'parsers se generaliza directamente a la producción de hipótesis: '
        'una gramática que describe la sintaxis de cadenas tipadas '
        'consistentes con la ontología (Σ_V, Σ_E) restringe al LLM a '
        'sugerir únicamente hipótesis bien-formadas, alimentadas por '
        'embeddings del grafo y scores GNN, complementando el catálogo '
        'ATT&CK existente.'
    )

    # ---------- Save ----------
    doc.save(str(SRC))
    print(f'wrote: {SRC}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
